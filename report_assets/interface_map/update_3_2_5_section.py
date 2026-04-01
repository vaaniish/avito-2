# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = r"C:\Study\diploma\avito-2\PZ_updated_3_2_draft_v12.docx"
OUT = r"C:\Study\diploma\avito-2\PZ_updated_3_2_draft_v13.docx"


def insert_after(paragraph, text, style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


doc = Document(SRC)

start = None
end = None
for i, p in enumerate(doc.paragraphs):
    t = (p.text or '').strip()
    if t == 'Определение стилистики' or t.startswith('3.2.5 '):
        start = i
        break

if start is None:
    raise RuntimeError('Не найден заголовок 3.2.5')

for i in range(start + 1, len(doc.paragraphs)):
    t = (doc.paragraphs[i].text or '').strip()
    if t == 'Дизайн-концепция' or t.startswith('3.2.6 '):
        end = i
        break

if end is None:
    raise RuntimeError('Не найден заголовок следующего подраздела после 3.2.5')

# Normalize headings
if doc.paragraphs[start].text.strip() != '3.2.5 Определение стилистики':
    doc.paragraphs[start].text = '3.2.5 Определение стилистики'

# Remove old body
for i in range(end - 1, start, -1):
    p = doc.paragraphs[i]
    p._element.getparent().remove(p._element)

anchor = doc.paragraphs[start]
style_normal = anchor.style

lines = [
    'Этап 5 «Определение стилистики» выполнялся после исследования и параллельно с проектированием экранов. Цель этапа — зафиксировать единое визуальное направление интерфейса до детальной проработки дизайн-концепции и исключить расхождения в стиле между пользовательским, партнерским и административным контурами.',
    'Для выбора стилистики были собраны несколько moodboard-наборов из референсных интерфейсов и UI-фрагментов: страницы маркетплейсов, кабинеты продавцов, примеры карточек товара, формы checkout, типографические и цветовые решения. В подборке использовались в первую очередь релевантные аналоги по предметной области (Авито, Яндекс Маркет, Ozon, Wildberries, seller-контуры), а не случайные визуальные примеры.',
    'Сравнение moodboard-наборов проводилось по прикладным критериям: читаемость на длительных сессиях, контрастность состояний, понятность иерархии блоков, стабильность форм и таблиц, а также визуальная совместимость с уже реализованными компонентами. По результатам выбран светлый утилитарный стиль: нейтральный фон, акцентный цвет для первичных действий, единая система отступов и радиусов, ограниченная палитра статусов (успех/ошибка/предупреждение).',
    'Для отчета рекомендуется использовать 2-3 скриншота референсной стилистики: 1) карточка и каталог маркетплейса; 2) checkout/формы; 3) кабинет продавца или администратора. Это наглядно показывает, на какой визуальный язык опирался выбор собственного стиля интерфейса.',
    'Сформированный style board (палитра, типографика, кнопки, поля ввода, таблицы и бейджи статусов) представлен на рисунке 22.',
    '[Вставить style board: 2-3 скриншота референсов + палитра/типографика/кнопки/поля/бейджи статусов]',
    'Рисунок 22 – Определение визуальной стилистики интерфейса',
]

for line in lines:
    anchor = insert_after(anchor, line, style=style_normal)

# Normalize next heading numbering
for i in range(start + 1, len(doc.paragraphs)):
    t = (doc.paragraphs[i].text or '').strip()
    if t == 'Дизайн-концепция':
        doc.paragraphs[i].text = '3.2.6 Дизайн-концепция'
        break

doc.save(OUT)
print(OUT)
