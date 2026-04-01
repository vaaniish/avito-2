# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = r"C:\Study\diploma\avito-2\PZ_updated_3_2_draft_v8.docx"
OUT = r"C:\Study\diploma\avito-2\PZ_updated_3_2_draft_v9.docx"

doc = Document(SRC)

start = None
end = None
for i, p in enumerate(doc.paragraphs):
    t = (p.text or '').strip()
    if t == 'Структура интерфейса' or t.startswith('3.2.3'):
        start = i
        break

if start is None:
    raise RuntimeError('Не найден заголовок 3.2.3')

for i in range(start + 1, len(doc.paragraphs)):
    t = (doc.paragraphs[i].text or '').strip()
    if t == 'Прототипирование интерфейса' or t.startswith('3.2.4'):
        end = i
        break

if end is None:
    raise RuntimeError('Не найден заголовок следующего подраздела после 3.2.3')

# Normalize heading text
if doc.paragraphs[start].text.strip() != '3.2.3 Структура интерфейса':
    doc.paragraphs[start].text = '3.2.3 Структура интерфейса'

# Remove old content between 3.2.3 and 3.2.4
for i in range(end - 1, start, -1):
    p = doc.paragraphs[i]
    p._element.getparent().remove(p._element)


def insert_after(paragraph, text, style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para

anchor = doc.paragraphs[start]
style_normal = anchor.style

new_lines = [
    'Этап 3 «Структура интерфейса» построен на результатах этапа пользовательских сценариев. Последовательности шагов из сценариев использованы как основа для формирования экранной карты: определены перечень экранов, их краткое содержание и положение в общей структуре интерфейса.',
    'В итоговую структуру включены 28 экранов: 12 экранов публичного контура, 5 экранов личного кабинета пользователя, 3 экрана партнерского контура и 8 экранов административного контура. Каждый экран представлен отдельным прямоугольным блоком, а связи между блоками отражают маршруты переходов и ролевые ограничения доступа.',
    'Такой способ представления структуры позволяет на раннем этапе проверить полноту покрытия сценариев, выявить избыточные переходы и зафиксировать точки входа в ключевые пользовательские потоки (каталог, checkout, профиль, партнерские операции и администрирование).',
    'Карта структуры экранов и маршрутов, где каждый экран вынесен в отдельный блок, представлена на рисунке 20.',
    '[Вставить изображение: report_assets/interface_map/interface_structure_screens.png]',
    'Рисунок 20 – Структура интерфейса и маршрутов системы',
]

for line in new_lines:
    anchor = insert_after(anchor, line, style=style_normal)

# Normalize next heading numbering (optional but useful for consistency)
for i in range(start + 1, len(doc.paragraphs)):
    t = (doc.paragraphs[i].text or '').strip()
    if t == 'Прототипирование интерфейса':
        doc.paragraphs[i].text = '3.2.4 Прототипирование интерфейса'
        break


doc.save(OUT)
print(OUT)
