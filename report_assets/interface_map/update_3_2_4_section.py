# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = r"C:\Study\diploma\avito-2\PZ_updated_3_2_draft_v9.docx"
OUT = r"C:\Study\diploma\avito-2\PZ_updated_3_2_draft_v10.docx"

doc = Document(SRC)

start = None
end = None
for i, p in enumerate(doc.paragraphs):
    t = (p.text or '').strip()
    if t == 'Прототипирование интерфейса' or t.startswith('3.2.4 '):
        start = i
        break

if start is None:
    raise RuntimeError('Не найден заголовок 3.2.4')

for i in range(start + 1, len(doc.paragraphs)):
    t = (doc.paragraphs[i].text or '').strip()
    if t == 'Определение стилистики' or t.startswith('3.2.5 '):
        end = i
        break

if end is None:
    raise RuntimeError('Не найден заголовок следующего подраздела после 3.2.4')

# Replace heading
if doc.paragraphs[start].text.strip() != '3.2.4 Прототипирование интерфейса':
    doc.paragraphs[start].text = '3.2.4 Прототипирование интерфейса'

# Remove old body of 3.2.4
for i in range(end - 1, start, -1):
    p = doc.paragraphs[i]
    p._element.getparent().remove(p._element)


def insert_after(paragraph, text, style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para

anchor = doc.paragraphs[start]
style_normal = anchor.style

new_lines = [
    'Этап 4 «Прототипирование интерфейса» выполнялся в двух итерациях: сначала формировался черновой (low-fidelity) прототип, затем уточненный финальный прототип. Такой подход позволил отделить проработку функциональной структуры экранов от визуального оформления и принимать решения на основе сценариев, а не на основе декоративных элементов интерфейса.',
    'Черновой прототип включал схематичные экраны и переходы между ними для основных потоков: каталог, карточка объявления, корзина и checkout, личный кабинет, партнерский и административный контуры. На этой итерации фиксировались зоны экранов, состав информационных блоков, порядок действий пользователя и точки вызова модальных окон без детальной стилизации.',
    'Финальный прототип использовался для проверки эргономики перед версткой: длины пользовательского пути, количества кликов до целевого действия, читаемости форм, полноты обязательных полей и согласованности состояний интерфейса. По результатам этапа были уточнены сетки, иерархия контента, расположение управляющих элементов и последовательность шагов в критичных сценариях.',
    'Черновой и финальный варианты прототипов ключевых экранов представлены на рисунке 21.',
    '[Вставить черновой и финальный прототипы экранов: каталог, карточка объявления, checkout, профиль, партнерский и административный кабинеты]',
    'Рисунок 21 – Прототипы ключевых экранов интерфейса',
]

for line in new_lines:
    anchor = insert_after(anchor, line, style=style_normal)

# Normalize next heading numbering
for i in range(start + 1, len(doc.paragraphs)):
    t = (doc.paragraphs[i].text or '').strip()
    if t == 'Определение стилистики':
        doc.paragraphs[i].text = '3.2.5 Определение стилистики'
        break

doc.save(OUT)
print(OUT)
