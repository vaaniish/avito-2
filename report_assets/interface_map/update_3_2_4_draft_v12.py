# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = r"C:\Study\diploma\avito-2\PZ_updated_3_2_draft_v10.docx"
OUT = r"C:\Study\diploma\avito-2\PZ_updated_3_2_draft_v12.docx"


def insert_after(paragraph, text, style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


doc = Document(SRC)

start = None
end = None
for i, p in enumerate(doc.paragraphs):
    t = (p.text or '').strip()
    if t.startswith('3.2.4 '):
        start = i
        break

if start is None:
    raise RuntimeError('Не найден заголовок 3.2.4')

for i in range(start + 1, len(doc.paragraphs)):
    t = (doc.paragraphs[i].text or '').strip()
    if t.startswith('3.2.5 '):
        end = i
        break

if end is None:
    raise RuntimeError('Не найден заголовок 3.2.5')

for i in range(end - 1, start, -1):
    p = doc.paragraphs[i]
    p._element.getparent().remove(p._element)

anchor = doc.paragraphs[start]
style_normal = anchor.style

lines = [
    'Этап 4 «Прототипирование интерфейса» выполнен в формате чернового (low-fidelity) прототипа. На данном этапе прорабатывались не визуальные детали, а функциональная структура экранов и логика переходов между ними.',
    'Черновой прототип включал схематичные экраны для ключевых потоков системы: каталог, карточка объявления, корзина и checkout, личный кабинет пользователя, партнерские разделы и административный контур. Для каждого экрана были зафиксированы основные зоны, состав блоков, порядок действий и точки вызова модальных форм.',
    'Результаты чернового прототипирования использовались для проверки объема интерфейса, длины пользовательского пути, количества переходов до целевого действия и достаточности полей ввода. По итогам этапа были уточнены сетки экранов, расположение функциональных блоков и структура маршрутов перед переходом к этапу стилистического оформления.',
    'Черновой прототип ключевых экранов интерфейса представлен на рисунке 21.',
    '[Вставить черновой прототип экранов: каталог, карточка объявления, checkout, профиль, партнерский и административный кабинеты]',
    'Рисунок 21 – Черновой прототип ключевых экранов интерфейса',
]

for line in lines:
    anchor = insert_after(anchor, line, style=style_normal)


doc.save(OUT)
print(OUT)
