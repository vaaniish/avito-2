# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import re

SRC = r"C:\Study\diploma\avito-2\PZ_updated_3_2_draft_v5.docx"
OUT = r"C:\Study\diploma\avito-2\PZ_updated_3_2_draft_v6.docx"

doc = Document(SRC)

caption_re = re.compile(r'^Рисунок\s+([0-9]+(?:\.[0-9]+)?)\s*[–-]\s*(.+)$')

# 1) Renumber all captions globally as Рисунок 1..N
caption_indices = []
for i, p in enumerate(doc.paragraphs):
    t = (p.text or '').strip()
    m = caption_re.match(t)
    if m:
        caption_indices.append((i, m.group(1), m.group(2)))

for new_num, (i, old_num, desc) in enumerate(caption_indices, start=1):
    doc.paragraphs[i].text = f'Рисунок {new_num} – {desc}'

# Rebuild lookup of caption number by paragraph index after renumbering
caption_number_by_index = {}
for i, p in enumerate(doc.paragraphs):
    t = (p.text or '').strip()
    m = re.match(r'^Рисунок\s+(\d+)\s*[–-]\s*(.+)$', t)
    if m:
        caption_number_by_index[i] = int(m.group(1))


def insert_before(paragraph, text):
    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = paragraph.style
    new_para.add_run(text)
    return new_para

# 2) Add references before 3.2.2 figure placeholders
placeholder_templates = {
    '[Вставить изображение: report_assets/scenarios/group_1_buyer.png]': 'Сценарии покупательского контура представлены на рисунке {n}.',
    '[Вставить изображение: report_assets/scenarios/group_2_profile.png]': 'Сценарии личного кабинета пользователя представлены на рисунке {n}.',
    '[Вставить изображение: report_assets/scenarios/group_3_partner.png]': 'Сценарии партнерского контура представлены на рисунке {n}.',
    '[Вставить изображение: report_assets/scenarios/group_4_admin.png]': 'Сценарии административного контура представлены на рисунке {n}.',
}

for i, p in enumerate(doc.paragraphs):
    text = (p.text or '').strip()
    if text not in placeholder_templates:
        continue

    # find nearest next caption
    next_caption_num = None
    for j in range(i + 1, len(doc.paragraphs)):
        if j in caption_number_by_index:
            next_caption_num = caption_number_by_index[j]
            break

    if next_caption_num is None:
        continue

    ref_text = placeholder_templates[text].format(n=next_caption_num)
    # avoid duplicate insertion if already present right above
    prev_text = (doc.paragraphs[i - 1].text or '').strip() if i > 0 else ''
    if prev_text != ref_text:
        insert_before(doc.paragraphs[i], ref_text)


doc.save(OUT)
print(OUT)
