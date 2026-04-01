# -*- coding: utf-8 -*-
from docx import Document
import re

SRC = r"C:\Study\diploma\avito-2\PZ_updated_3_2_draft_v7.docx"
OUT = r"C:\Study\diploma\avito-2\PZ_updated_3_2_draft_v8.docx"

doc = Document(SRC)

# caption indices -> number
caption_indices = []
for i, p in enumerate(doc.paragraphs):
    t = (p.text or '').strip()
    m = re.match(r'^Рисунок\s+(\d+)\s*[–-]\s*(.+)$', t)
    if m:
        caption_indices.append((i, int(m.group(1))))

if not caption_indices:
    doc.save(OUT)
    print(OUT)
    raise SystemExit

# helper: nearest next caption number for a paragraph index

def next_caption_num(idx):
    for ci, num in caption_indices:
        if ci > idx:
            return num
    return caption_indices[-1][1]

pattern = re.compile(r'(?i)(на\s+рисунке\s+)(\d+(?:\.\d+)?)')

for i, p in enumerate(doc.paragraphs):
    t = p.text or ''
    st = t.strip()

    # skip caption lines themselves
    if re.match(r'^Рисунок\s+\d+\s*[–-]\s*', st):
        continue

    if not pattern.search(t):
        continue

    target_num = next_caption_num(i)

    def repl(m):
        return f"{m.group(1)}{target_num}"

    new_t = pattern.sub(repl, t)
    if new_t != t:
        p.text = new_t


doc.save(OUT)
print(OUT)
