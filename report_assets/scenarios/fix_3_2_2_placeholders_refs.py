# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = r"C:\Study\diploma\avito-2\PZ_updated_3_2_draft_v6.docx"
OUT = r"C:\Study\diploma\avito-2\PZ_updated_3_2_draft_v7.docx"

doc = Document(SRC)


def insert_before(paragraph, text):
    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = paragraph.style
    new_para.add_run(text)
    return new_para

cases = [
    {
        'caption_contains': 'Покупательский контур',
        'placeholder': '[Вставить изображение: report_assets/scenarios/group_1_buyer.png]',
        'ref': 'Сценарии покупательского контура представлены на рисунке 16.'
    },
    {
        'caption_contains': 'Личный кабинет пользователя',
        'placeholder': '[Вставить изображение: report_assets/scenarios/group_2_profile.png]',
        'ref': 'Сценарии личного кабинета пользователя представлены на рисунке 17.'
    },
    {
        'caption_contains': 'Партнерский контур',
        'placeholder': '[Вставить изображение: report_assets/scenarios/group_3_partner.png]',
        'ref': 'Сценарии партнерского контура представлены на рисунке 18.'
    },
    {
        'caption_contains': 'Административный контур',
        'placeholder': '[Вставить изображение: report_assets/scenarios/group_4_admin.png]',
        'ref': 'Сценарии административного контура представлены на рисунке 19.'
    },
]

for c in cases:
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or '').strip()
        if t.startswith('Рисунок ') and c['caption_contains'] in t:
            # ensure placeholder paragraph exists right above caption
            prev = doc.paragraphs[i - 1]
            prev.text = c['placeholder']

            # ensure reference sentence before placeholder
            prev2_text = (doc.paragraphs[i - 2].text or '').strip() if i >= 2 else ''
            if prev2_text != c['ref']:
                insert_before(prev, c['ref'])
            break


doc.save(OUT)
print(OUT)
