from pathlib import Path

import docx


FILES = [
    Path(r"c:\Study\diploma\avito-2\PZ_updated_3_1_fixed.docx"),
    Path(r"C:\Study\Методы проектирования\Доки для курсовой\PZ_updated_3_1_fixed.docx"),
]

INTRO_TEXT = (
    "В разделе рассматриваются результаты реализации программной части "
    "информационной системы: структура серверных модулей, ключевые "
    "бизнес-процессы, правила обработки пользовательских и финансовых "
    "операций, а также принятые решения по обеспечению целостности данных "
    "и устойчивости работы приложения."
)


def process_file(path: Path) -> None:
    if not path.exists():
        print(f"SKIP (not found): {path}")
        return

    doc = docx.Document(str(path))

    # Находим раздел 3 и подпункт 3.1 как последние заголовки перед "Заключением".
    h1 = [i for i, p in enumerate(doc.paragraphs) if p.style and p.style.name == "1"]
    if len(h1) < 2:
        print(f"SKIP (no level-1 headings): {path}")
        return
    idx_conclusion = h1[-1]

    h2 = [
        i
        for i, p in enumerate(doc.paragraphs[:idx_conclusion])
        if p.style and p.style.name == "2"
    ]
    if len(h2) < 2:
        print(f"SKIP (no level-2 headings): {path}")
        return
    idx_31 = h2[-2]

    idx_3 = None
    for i in range(idx_31 - 1, -1, -1):
        p = doc.paragraphs[i]
        if p.style and p.style.name == "1":
            idx_3 = i
            break
    if idx_3 is None:
        print(f"SKIP (section 3 heading not found): {path}")
        return

    # Удаляем все непустые абзацы между разделом 3 и 3.1 (чтобы убрать испорченный текст).
    i = idx_3 + 1
    while i < idx_31:
        p = doc.paragraphs[i]
        if p.text.strip():
            p._element.getparent().remove(p._element)
            idx_31 -= 1
            continue
        i += 1

    # Определяем стиль основного текста.
    body_style = None
    for j in range(idx_31 + 1, len(doc.paragraphs)):
        p = doc.paragraphs[j]
        if p.text.strip() and p.style and p.style.name not in {"1", "2", "3"}:
            body_style = p.style
            break
    if body_style is None:
        body_style = doc.styles["Normal"]

    anchor = doc.paragraphs[idx_31]
    intro = anchor.insert_paragraph_before(INTRO_TEXT)
    intro.style = body_style

    doc.save(str(path))
    print(f"UPDATED: {path}")


def main() -> None:
    for file_path in FILES:
        process_file(file_path)


if __name__ == "__main__":
    main()
