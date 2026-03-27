from __future__ import annotations

import re
from pathlib import Path

import docx


DOC_PATH = Path(r"c:\Study\diploma\avito-2\PZ_updated_3_1.docx")
MIGRATION_PATH = Path(
    r"c:\Study\diploma\avito-2\backend\prisma\migrations\20260323170000_init_squashed\migration.sql"
)

SECTION_PARAGRAPHS = [
    (
        "В данном подразделе описана реализация бизнес-логики серверной части "
        "информационной системы электронной торговой площадки и реализованные "
        "ограничения целостности данных. Серверная часть построена по модульному "
        "принципу и включает подсистемы аутентификации, каталога, профиля, "
        "партнерского кабинета и администрирования."
    ),
    "Структура backend-модулей и их взаимодействие с внешними сервисами и слоем доступа к данным представлены на рисунке 3.1.",
    "[Место для вставки скриншота структуры backend-модулей]",
    "Рисунок 3.1 – Структура backend-модулей серверной части",
    (
        "Разделение ответственности реализовано следующим образом: маршруты API "
        "сгруппированы по бизнес-доменам, а операции чтения и записи централизованы "
        "через prisma layer. Для контроля доступа применяются серверные проверки "
        "ролей и статуса пользователя (покупатель, партнер, администратор), включая "
        "блокировку выполнения операций для пользователей со статусом BLOCKED."
    ),
    (
        "В партнерском модуле реализован управляемый жизненный цикл объявления: "
        "при создании и редактировании выполняются проверки обязательных полей, "
        "объявление переводится в модерацию, а прямой перевод в активное состояние "
        "со стороны партнера запрещен. В заказах реализовано формирование заказов "
        "по продавцам, расчет комиссии и создание финансовых транзакций в рамках "
        "единой транзакции БД."
    ),
    (
        "В контуре доставки реализованы проверки трек-номеров и ограничения смены "
        "статусов, а в контуре жалоб – антидублирование, ограничение частоты подачи "
        "жалоб, регистрация событий и идемпотентная обработка административного "
        "изменения статуса жалобы. При подтверждении жалобы транзакционно "
        "применяются санкции к продавцу и связанные изменения состояния объявления."
    ),
    (
        "Для пользовательских данных реализованы дополнительные ограничения: "
        "поддержка только одного адреса по умолчанию, запрет удаления адреса "
        "по умолчанию, идемпотентное добавление в избранное, а также проверка "
        "условий публикации отзывов (только после завершенной покупки, без "
        "повторного отзыва от того же пользователя для того же товара)."
    ),
    (
        "Систематизированный перечень ограничений целостности данных приведен в "
        "таблице 3.1. Полный SQL-скрипт реализации ограничений представлен в "
        "Приложении Б."
    ),
]


def parse_migration_constraints(sql_text: str) -> list[tuple[str, str, str, str]]:
    rows: list[tuple[str, str, str, str]] = []
    lines = sql_text.splitlines()

    unique_re = re.compile(
        r'^CREATE\s+UNIQUE\s+INDEX\s+"(?P<name>[^"]+)"\s+ON\s+"(?P<table>[^"]+)"\s*\((?P<cols>[^)]*)\)\s*(?P<rest>.*)$',
        re.IGNORECASE,
    )
    fk_re = re.compile(
        r'^ALTER\s+TABLE\s+"(?P<table>[^"]+)"\s+ADD\s+CONSTRAINT\s+"(?P<name>[^"]+)"\s+FOREIGN\s+KEY\s*\((?P<cols>[^)]*)\)\s+REFERENCES\s+"(?P<ref_table>[^"]+)"\s*\((?P<ref_cols>[^)]*)\)\s+ON\s+DELETE\s+(?P<on_delete>\w+)\s+ON\s+UPDATE\s+(?P<on_update>\w+)$',
        re.IGNORECASE,
    )
    alter_table_re = re.compile(r'^ALTER\s+TABLE\s+"(?P<table>[^"]+)"$', re.IGNORECASE)
    add_chk_re = re.compile(
        r'^ADD\s+CONSTRAINT\s+"(?P<name>[^"]+_chk)"$',
        re.IGNORECASE,
    )
    check_expr_re = re.compile(r"^CHECK\s*\((?P<expr>.*)\)\s*,?$", re.IGNORECASE)

    current_table: str | None = None
    pending_check_name: str | None = None
    pending_check_table: str | None = None

    for raw_line in lines:
        line = raw_line.strip()
        if not line or line.startswith("--"):
            continue
        line = line.rstrip(";").strip()

        m_unique = unique_re.match(line)
        if m_unique:
            cols = m_unique.group("cols").replace('"', "")
            rest = m_unique.group("rest").strip().strip(";").strip()
            restriction = cols if not rest else f"{cols}; {rest}"
            rows.append(
                (
                    "UNIQUE INDEX",
                    m_unique.group("table"),
                    restriction,
                    m_unique.group("name"),
                )
            )
            continue

        m_fk = fk_re.match(line)
        if m_fk:
            cols = m_fk.group("cols").replace('"', "")
            ref_cols = m_fk.group("ref_cols").replace('"', "")
            restriction = (
                f"({cols}) -> {m_fk.group('ref_table')}({ref_cols}), "
                f"ON DELETE {m_fk.group('on_delete')}, ON UPDATE {m_fk.group('on_update')}"
            )
            rows.append(
                (
                    "FOREIGN KEY",
                    m_fk.group("table"),
                    restriction,
                    m_fk.group("name"),
                )
            )
            continue

        m_alter = alter_table_re.match(line)
        if m_alter:
            current_table = m_alter.group("table")
            pending_check_name = None
            pending_check_table = None
            continue

        m_add_chk = add_chk_re.match(line.rstrip(","))
        if m_add_chk and current_table:
            pending_check_name = m_add_chk.group("name")
            pending_check_table = current_table
            continue

        if pending_check_name and pending_check_table:
            m_expr = check_expr_re.match(line)
            if m_expr:
                rows.append(
                    (
                        "CHECK",
                        pending_check_table,
                        m_expr.group("expr"),
                        pending_check_name,
                    )
                )
                pending_check_name = None
                pending_check_table = None
                continue

    # Исключаем устаревшие ограничения для удаленной сущности City.
    filtered: list[tuple[str, str, str, str]] = []
    for r in rows:
        type_name, table_name, restriction, obj_name = r
        text = f"{table_name} {restriction} {obj_name}".lower()
        if table_name == "City":
            continue
        if "city_id" in text:
            continue
        filtered.append((type_name, table_name, restriction, obj_name))

    # Дедупликация по имени ограничения/индекса.
    seen: set[str] = set()
    unique_rows: list[tuple[str, str, str, str]] = []
    for row in filtered:
        key = row[3]
        if key in seen:
            continue
        seen.add(key)
        unique_rows.append(row)

    return unique_rows


def main() -> None:
    if not DOC_PATH.exists():
        raise FileNotFoundError(f"Не найден файл отчета: {DOC_PATH}")
    if not MIGRATION_PATH.exists():
        raise FileNotFoundError(f"Не найден migration.sql: {MIGRATION_PATH}")

    doc = docx.Document(str(DOC_PATH))

    heading1_idxs = [
        i for i, p in enumerate(doc.paragraphs) if p.style and p.style.name == "1"
    ]
    if len(heading1_idxs) < 2:
        raise RuntimeError("Не найдены основные разделы (стиль 1)")
    idx_conclusion = heading1_idxs[-1]

    level2_before_conclusion = [
        i
        for i, p in enumerate(doc.paragraphs[:idx_conclusion])
        if p.style and p.style.name == "2"
    ]
    if len(level2_before_conclusion) < 2:
        raise RuntimeError("Не найдены подразделы 3.1/3.2 (стиль 2)")

    idx_31 = level2_before_conclusion[-2]
    idx_32 = level2_before_conclusion[-1]
    if idx_32 <= idx_31:
        raise RuntimeError("Нарушен порядок 3.1/3.2")

    # Найти стиль основного текста (обычно "Деф текст")
    body_style = None
    for j in range(idx_31 - 1, -1, -1):
        p = doc.paragraphs[j]
        if p.text.strip() and p.style and p.style.name not in {"1", "2", "3"}:
            body_style = p.style
            break
    if body_style is None:
        body_style = doc.styles["Normal"]

    # Очистить текущее содержание между 3.1 и 3.2
    for _ in range(idx_32 - idx_31 - 1):
        p = doc.paragraphs[idx_31 + 1]
        p._element.getparent().remove(p._element)

    # Вставить новый текст 3.1 перед заголовком 3.2
    anchor = doc.paragraphs[idx_31 + 1]
    for text in SECTION_PARAGRAPHS:
        p = anchor.insert_paragraph_before(text)
        p.style = body_style

    # Добавить подпись таблицы
    p = anchor.insert_paragraph_before(
        "Таблица 3.1 – Ограничения целостности данных информационной системы"
    )
    p.style = body_style

    # Построить таблицу ограничений из migration.sql
    sql_text = MIGRATION_PATH.read_text(encoding="utf-8", errors="replace")
    rows = parse_migration_constraints(sql_text)
    if not rows:
        raise RuntimeError("Не удалось извлечь ограничения из migration.sql")

    table = doc.add_table(rows=1, cols=4)
    header = table.rows[0].cells
    header[0].text = "Тип ограничения"
    header[1].text = "Таблица"
    header[2].text = "Ограничение"
    header[3].text = "Имя ограничения/индекса"

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]
        cells[2].text = row[2]
        cells[3].text = row[3]

    # Переместить таблицу в место 3.1 (перед 3.2)
    anchor._p.addprevious(table._tbl)

    p = anchor.insert_paragraph_before(
        "Использование перечисленных ограничений обеспечивает непротиворечивость данных при многопользовательской работе, корректную обработку финансовых операций и устойчивость бизнес-процессов к ошибочным и конкурентным запросам."
    )
    p.style = body_style

    doc.save(str(DOC_PATH))
    print(f"OK: раздел 3.1 обновлен, ограничений в таблице: {len(rows)}")


if __name__ == "__main__":
    main()
